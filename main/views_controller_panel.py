# Importlarni yuqoriga ko'chirish
from main.models import GroupSubject, Semester, Group, Bulim, Kafedra, Subject, University, Faculty
# AJAX orqali guruhga tegishli fanlarni qaytaruvchi endpoint
from django.views.decorators.http import require_GET
@require_GET
def get_subjects_by_group(request):
    group_id = request.GET.get('group_id')
    subjects = []
    if group_id:
        group_subjects = GroupSubject.objects.filter(group_id=group_id)
        for gs in group_subjects.select_related('subject', 'semester'):
            subjects.append({
                'id': gs.subject.id,
                'name': gs.subject.name,
                'semester': gs.semester.number if gs.semester else None
            })
    return JsonResponse({'subjects': subjects})
from collections import defaultdict
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import io
import openpyxl
from django.http import HttpResponse
from main.models import User, Kafedra, Bulim
from django.views.decorators.http import require_POST
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.shortcuts import render, redirect
from main.models import Test, Subject, Question, Group
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required, user_passes_test
# Faqat controller uchun dekorator
def controller_required(view_func):
    return user_passes_test(lambda u: u.is_authenticated and u.role == 'controller')(view_func)

# --- WORD EXPORT ---

# GroupSubject ro'yxati (faqat controller)
@controller_required


@controller_required
def group_subjects_list(request):

    # Universitet o'chirish
    delete_university_id = request.GET.get('delete_university')
    if delete_university_id:
        try:
            University.objects.filter(id=delete_university_id).delete()
            msg = 'Universitet o‘chirildi.'
        except Exception as e:
            msg = f'Universitet o‘chirishda xatolik: {e}'

    # Universitet tahrirlash
    edit_university_id = request.GET.get('edit_university')
    edit_university = None
    if edit_university_id:
        try:
            edit_university = University.objects.get(id=edit_university_id)
        except University.DoesNotExist:
            edit_university = None
            msg = 'Universitet topilmadi.'
        if request.method == 'POST' and request.POST.get('edit_university_id') == edit_university_id:
            edit_university.name = request.POST.get('university_name')
            try:
                edit_university.full_clean()
                edit_university.save()
                msg = 'Universitet tahrirlandi.'
                edit_university = None
            except Exception as e:
                msg = f'Universitet tahrirda xatolik: {e}'

    # Fakultet o'chirish
    delete_faculty_id = request.GET.get('delete_faculty')
    if delete_faculty_id:
        try:
            Faculty.objects.filter(id=delete_faculty_id).delete()
            msg = 'Fakultet o‘chirildi.'
        except Exception as e:
            msg = f'Fakultet o‘chirishda xatolik: {e}'

    # Fakultet tahrirlash
    edit_faculty_id = request.GET.get('edit_faculty')
    edit_faculty = None
    if edit_faculty_id:
        try:
            edit_faculty = Faculty.objects.get(id=edit_faculty_id)
        except Faculty.DoesNotExist:
            edit_faculty = None
            msg = 'Fakultet topilmadi.'
        if request.method == 'POST' and request.POST.get('edit_faculty_id') == edit_faculty_id:
            edit_faculty.name = request.POST.get('faculty_name')
            edit_faculty.university_id = request.POST.get('university')
            try:
                edit_faculty.full_clean()
                edit_faculty.save()
                msg = 'Fakultet tahrirlandi.'
                edit_faculty = None
            except Exception as e:
                msg = f'Fakultet tahrirda xatolik: {e}'

    # Bo'lim o'chirish
    delete_bulim_id = request.GET.get('delete_bulim')
    if delete_bulim_id:
        try:
            Bulim.objects.filter(id=delete_bulim_id).delete()
            msg = 'Bo‘lim o‘chirildi.'
        except Exception as e:
            msg = f'Bo‘lim o‘chirishda xatolik: {e}'

    # Bo'lim tahrirlash
    edit_bulim_id = request.GET.get('edit_bulim')
    edit_bulim = None
    if edit_bulim_id:
        try:
            edit_bulim = Bulim.objects.get(id=edit_bulim_id)
        except Bulim.DoesNotExist:
            edit_bulim = None
            msg = 'Bo‘lim topilmadi.'
        if request.method == 'POST' and request.POST.get('edit_bulim_id') == edit_bulim_id:
            edit_bulim.name = request.POST.get('bulim_name')
            try:
                edit_bulim.full_clean()
                edit_bulim.save()
                msg = 'Bo‘lim tahrirlandi.'
                edit_bulim = None
            except Exception as e:
                msg = f'Bo‘lim tahrirda xatolik: {e}'

    # Guruh o'chirish
    delete_group_id = request.GET.get('delete_group')
    if delete_group_id:
        try:
            Group.objects.filter(id=delete_group_id).delete()
            msg = 'Guruh o‘chirildi.'
        except Exception as e:
            msg = f'Guruh o‘chirishda xatolik: {e}'

    # Guruh tahrirlash
    edit_group_id = request.GET.get('edit_group')
    edit_group = None
    if edit_group_id:
        try:
            edit_group = Group.objects.get(id=edit_group_id)
        except Group.DoesNotExist:
            edit_group = None
            msg = 'Guruh topilmadi.'
        if request.method == 'POST' and request.POST.get('edit_group_id') == edit_group_id:
            edit_group.name = request.POST.get('group_name')
            edit_group.faculty_id = request.POST.get('faculty')
            try:
                edit_group.full_clean()
                edit_group.save()
                msg = 'Guruh tahrirlandi.'
                edit_group = None
            except Exception as e:
                msg = f'Guruh tahrirda xatolik: {e}'

    # Fan o'chirish
    delete_subject_id = request.GET.get('delete_subject')
    if delete_subject_id:
        try:
            Subject.objects.filter(id=delete_subject_id).delete()
            msg = 'Fan o‘chirildi.'
        except Exception as e:
            msg = f'Fan o‘chirishda xatolik: {e}'

    # Fan tahrirlash
    edit_subject_id = request.GET.get('edit_subject')
    edit_subject = None
    if edit_subject_id:
        try:
            edit_subject = Subject.objects.get(id=edit_subject_id)
        except Subject.DoesNotExist:
            edit_subject = None
            msg = 'Fan topilmadi.'
        if request.method == 'POST' and request.POST.get('edit_subject_id') == edit_subject_id:
            edit_subject.name = request.POST.get('subject_name')
            try:
                edit_subject.full_clean()
                edit_subject.save()
                msg = 'Fan tahrirlandi.'
                edit_subject = None
            except Exception as e:
                msg = f'Fan tahrirda xatolik: {e}'

    # Kafedra o'chirish
    delete_kafedra_id = request.GET.get('delete_kafedra')
    if delete_kafedra_id:
        try:
            Kafedra.objects.filter(id=delete_kafedra_id).delete()
            msg = 'Kafedra o‘chirildi.'
        except Exception as e:
            msg = f'Kafedra o‘chirishda xatolik: {e}'

    # Kafedra tahrirlash (GET: formni to'ldirish, POST: saqlash)
    edit_kafedra_id = request.GET.get('edit_kafedra')
    edit_kafedra = None
    if edit_kafedra_id:
        try:
            edit_kafedra = Kafedra.objects.get(id=edit_kafedra_id)
        except Kafedra.DoesNotExist:
            edit_kafedra = None
            msg = 'Kafedra topilmadi.'
        # Tahrirlashni saqlash
        if request.method == 'POST' and request.POST.get('edit_kafedra_id') == edit_kafedra_id:
            edit_kafedra.name = request.POST.get('kafedra_name')
            edit_kafedra.faculty_id = request.POST.get('faculty')
            try:
                edit_kafedra.full_clean()
                edit_kafedra.save()
                msg = 'Kafedra tahrirlandi.'
                edit_kafedra = None
            except Exception as e:
                msg = f'Kafedra tahrirda xatolik: {e}'
    msg = ''
    edit_gs = None

    # O'chirish
    delete_id = request.GET.get('delete')
    if delete_id:
        try:
            GroupSubject.objects.filter(id=delete_id).delete()
            msg = 'Bog‘lash o‘chirildi.'
        except Exception as e:
            msg = f'O‘chirishda xatolik: {e}'



    # Yangi universitet qo'shish
    if request.method == 'POST' and 'add_university' in request.GET:
        name = request.POST.get('university_name')
        if name:
            University.objects.create(name=name)
            msg = 'Universitet qo‘shildi.'

    # Yangi fakultet qo'shish
    if request.method == 'POST' and 'add_faculty' in request.GET:
        name = request.POST.get('faculty_name')
        university_id = request.POST.get('university')
        if name and university_id:
            Faculty.objects.create(name=name, university_id=university_id)
            msg = 'Fakultet qo‘shildi.'

    # Yangi kafedra qo'shish (fakultetdan keyin)
    if request.method == 'POST' and 'add_kafedra' in request.GET:
        name = request.POST.get('kafedra_name')
        faculty_id = request.POST.get('faculty')
        if name and faculty_id:
            Kafedra.objects.create(name=name, faculty_id=faculty_id)
            msg = 'Kafedra qo‘shildi.'

    # Yangi bo'lim qo'shish
    if request.method == 'POST' and 'add_bulim' in request.GET:
        name = request.POST.get('bulim_name')
        if name:
            Bulim.objects.create(name=name)
            msg = 'Bo‘lim qo‘shildi.'

    # Yangi guruh qo'shish
    if request.method == 'POST' and 'add_group' in request.GET:
        name = request.POST.get('group_name')
        faculty_id = request.POST.get('faculty')
        if name and faculty_id:
            Group.objects.create(name=name, faculty_id=faculty_id)
            msg = 'Guruh qo‘shildi.'

    # Yangi fan qo'shish
    if request.method == 'POST' and 'add_subject' in request.GET:
        name = request.POST.get('subject_name')
        if name:
            Subject.objects.create(name=name)
            msg = 'Fan qo‘shildi.'

    # Tahrirlash (GET: formni to'ldirish, POST: saqlash)
    edit_id = request.GET.get('edit')
    if edit_id:
        try:
            edit_gs = GroupSubject.objects.get(id=edit_id)
        except GroupSubject.DoesNotExist:
            edit_gs = None
            msg = 'Bog‘lash topilmadi.'
        # Tahrirlashni saqlash
        if request.method == 'POST' and request.POST.get('edit_id') == edit_id:
            edit_gs.group_id = request.POST.get('group') or None
            edit_gs.bulim_id = request.POST.get('bulim') or None
            edit_gs.kafedra_id = request.POST.get('kafedra') or None
            edit_gs.subject_id = request.POST.get('subject')
            edit_gs.semester_id = request.POST.get('semester') or None
            try:
                edit_gs.full_clean()
                edit_gs.save()
                msg = 'Tahrirlandi.'
                edit_gs = None
            except Exception as e:
                msg = f'Tahrirda xatolik: {e}'
    # GroupSubject qo'shish
    elif request.method == 'POST':
        group_id = request.POST.get('group') or None
        bulim_id = request.POST.get('bulim') or None
        kafedra_id = request.POST.get('kafedra') or None
        subject_id = request.POST.get('subject')
        semester_id = request.POST.get('semester') or None
        gs = GroupSubject(
            group_id=group_id if group_id else None,
            bulim_id=bulim_id if bulim_id else None,
            kafedra_id=kafedra_id if kafedra_id else None,
            subject_id=subject_id,
            semester_id=semester_id if semester_id else None
        )
        try:
            gs.full_clean()
            gs.save()
            msg = 'Bog‘lash muvaffaqiyatli qo‘shildi.'
        except Exception as e:
            msg = f'Xatolik: {e}'

    group_subjects = GroupSubject.objects.select_related('group', 'bulim', 'kafedra', 'subject', 'semester').all()

    context = {
        'group_subjects': group_subjects,
        'groups': Group.objects.all(),
        'bulimlar': Bulim.objects.all(),
        'kafedralar': Kafedra.objects.all(),
        'subjects': Subject.objects.all(),
        'semesters': Semester.objects.all(),
        'universities': University.objects.all(),
        'faculties': Faculty.objects.all(),
        'msg': msg,
        'edit_gs': edit_gs,
        'edit_kafedra': edit_kafedra,
        'edit_university': edit_university,
        'edit_faculty': edit_faculty,
        'edit_bulim': edit_bulim,
        'edit_group': edit_group,
        'edit_subject': edit_subject,
    }
    return render(request, 'controller_panel/group_subjects.html', context)
from django.utils import timezone

# --- WORD EXPORT ---
@login_required
def export_users_word(request):
    users = User.objects.exclude(is_superuser=True)
    # Filterlar
    filter_role = request.GET.get('filter_role')
    filter_group = request.GET.get('filter_group')
    filter_kafedra = request.GET.get('filter_kafedra')
    filter_bulim = request.GET.get('filter_bulim')
    if filter_role:
        users = users.filter(role=filter_role)
    if filter_group:
        users = users.filter(group_id=filter_group)
    if filter_kafedra:
        users = users.filter(kafedra_id=filter_kafedra)
    if filter_bulim:
        users = users.filter(bulim_id=filter_bulim)
    # Role bo‘yicha ajratish
    role_map = defaultdict(list)
    for user in users:
        role_map[user.get_role_display()].append(user)
    doc = Document()
    doc.add_heading('Foydalanuvchilar ro‘yxati', 0)
    for role, userlist in role_map.items():
        doc.add_heading(role, level=1)
        table = doc.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Ism'
        hdr_cells[1].text = 'Familiya'
        hdr_cells[2].text = 'Username'
        hdr_cells[3].text = 'Guruh/Kafedra/Bo‘lim'
        hdr_cells[4].text = 'Access code'
        hdr_cells[5].text = 'Role'
        seen = set()
        for u in userlist:
            key = (u.first_name, u.last_name, u.username, u.role, u.group_id, u.kafedra_id, u.bulim_id)
            if key in seen:
                continue
            seen.add(key)
            row = table.add_row().cells
            row[0].text = u.first_name or ''
            row[1].text = u.last_name or ''
            row[2].text = u.username or ''
            # Guruh/kafedra/bulim bir ustunda
            if u.role == 'student' and u.group:
                row[3].text = u.group.name
            elif u.role == 'tutor' and u.kafedra:
                row[3].text = u.kafedra.name
            elif u.role == 'employee' and u.bulim:
                row[3].text = u.bulim.name
            else:
                row[3].text = ''
            row[4].text = u.access_code or ''
            row[5].text = u.get_role_display()
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    response = HttpResponse(output.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=foydalanuvchilar.docx'
    return response

# --- PDF EXPORT ---
@login_required
def export_users_pdf(request):
    users = User.objects.exclude(is_superuser=True)
    filter_role = request.GET.get('filter_role')
    filter_group = request.GET.get('filter_group')
    filter_kafedra = request.GET.get('filter_kafedra')
    filter_bulim = request.GET.get('filter_bulim')
    if filter_role:
        users = users.filter(role=filter_role)
    if filter_group:
        users = users.filter(group_id=filter_group)
    if filter_kafedra:
        users = users.filter(kafedra_id=filter_kafedra)
    if filter_bulim:
        users = users.filter(bulim_id=filter_bulim)
    role_map = defaultdict(list)
    for user in users:
        role_map[user.get_role_display()].append(user)
    output = io.BytesIO()
    c = canvas.Canvas(output, pagesize=A4)
    width, height = A4
    y = height - 40
    c.setFont('Helvetica-Bold', 16)
    c.drawString(40, y, 'Foydalanuvchilar ro‘yxati')
    y -= 30
    c.setFont('Helvetica', 12)
    for role, userlist in role_map.items():
        c.setFont('Helvetica-Bold', 14)
        c.drawString(40, y, role)
        y -= 22
        c.setFont('Helvetica', 10)
        c.drawString(40, y, 'Ism')
        c.drawString(120, y, 'Familiya')
        c.drawString(220, y, 'Username')
        c.drawString(320, y, 'Guruh/Kafedra/Bo‘lim')
        c.drawString(470, y, 'Access code')
        y -= 16
        seen = set()
        for u in userlist:
            key = (u.first_name, u.last_name, u.username, u.role, u.group_id, u.kafedra_id, u.bulim_id)
            if key in seen:
                continue
            seen.add(key)
            c.drawString(40, y, u.first_name or '')
            c.drawString(120, y, u.last_name or '')
            c.drawString(220, y, u.username or '')
            if u.role == 'student' and u.group:
                c.drawString(320, y, u.group.name)
            elif u.role == 'tutor' and u.kafedra:
                c.drawString(320, y, u.kafedra.name)
            elif u.role == 'employee' and u.bulim:
                c.drawString(320, y, u.bulim.name)
            else:
                c.drawString(320, y, '')
            c.drawString(470, y, u.access_code or '')
            y -= 14
            if y < 60:
                c.showPage()
                y = height - 40
                c.setFont('Helvetica', 10)
        y -= 18
    c.save()
    output.seek(0)
    response = HttpResponse(output, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename=foydalanuvchilar.pdf'
    return response
@login_required
def export_users_excel(request):
    filter_role = request.GET.get('filter_role')
    users = User.objects.exclude(is_superuser=True)
    if filter_role:
        users = users.filter(role=filter_role)
    filter_group = request.GET.get('filter_group')
    filter_kafedra = request.GET.get('filter_kafedra')
    filter_bulim = request.GET.get('filter_bulim')
    if filter_group:
        users = users.filter(group_id=filter_group)
    if filter_kafedra:
        users = users.filter(kafedra_id=filter_kafedra)
    if filter_bulim:
        users = users.filter(bulim_id=filter_bulim)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Foydalanuvchilar'
    ws.append(['Ism', 'Familiya', 'Access code'])
    for user in users:
        ws.append([
            getattr(user, 'first_name', ''),
            getattr(user, 'last_name', ''),
            getattr(user, 'access_code', '')
        ])
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    response = HttpResponse(output, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=foydalanuvchilar.xlsx'
    return response


# Foydalanuvchilarni ko‘rish va qo‘shish (superuserlarsiz)
@login_required
def add_user(request):
    if not hasattr(request.user, 'role') or request.user.role != 'controller':
        return redirect('/api/login/')
    users = User.objects.exclude(is_superuser=True)
    # Filter GET params
    filter_role = request.GET.get('filter_role')
    filter_group = request.GET.get('filter_group')
    filter_kafedra = request.GET.get('filter_kafedra')
    filter_bulim = request.GET.get('filter_bulim')
    if filter_role:
        users = users.filter(role=filter_role)
    if filter_group:
        users = users.filter(group_id=filter_group)
    if filter_kafedra:
        users = users.filter(kafedra_id=filter_kafedra)
    if filter_bulim:
        users = users.filter(bulim_id=filter_bulim)

    groups = Group.objects.all()
    kafedralar = Kafedra.objects.all()
    bulimlar = Bulim.objects.all()
    role_choices = User.ROLE_CHOICES
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        role = request.POST.get('role')
        group_id = request.POST.get('group')
        kafedra_id = request.POST.get('kafedra')
        bulim_id = request.POST.get('bulim')
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        user = User(username=username, role=role, first_name=first_name, last_name=last_name)
        if group_id:
            user.group_id = group_id
        if kafedra_id:
            user.kafedra_id = kafedra_id
        if bulim_id:
            user.bulim_id = bulim_id
        user.set_password(password)
        user.save()
        return redirect('add_user')
    return render(request, 'controller_panel/add_user.html', {
        'users': users,
        'groups': groups,
        'kafedralar': kafedralar,
        'bulimlar': bulimlar,
        'role_choices': role_choices,
    })
@login_required
def controller_logout(request):
    logout(request)
    return redirect('/api/login/')
# AJAX: Controller o‘zining savolini o‘chira oladi (lekin tahrirlay olmaydi)

def login_check(request):
    if not request.user.is_authenticated:
        return redirect('/api/login/')
    return None

@login_required
@require_POST
def delete_question(request, question_id):
    login_redirect = login_check(request)
    if login_redirect:
        return login_redirect
    if not hasattr(request.user, 'role') or request.user.role != 'controller':
        return JsonResponse({'success': False, 'error': 'Ruxsat yo‘q'}, status=403)
    deleted, _ = Question.objects.filter(id=question_id).delete()
    return JsonResponse({'success': bool(deleted)})


# AJAX: Fanga tegishli savollar ro‘yxati (faqat controller o‘zining savollarini ko‘radi)
from django.views.decorators.http import require_GET
@login_required
@require_GET
def subject_questions(request, subject_id):
    login_redirect = login_check(request)
    if login_redirect:
        return login_redirect
    if not hasattr(request.user, 'role') or request.user.role != 'controller':
        return JsonResponse({'error': 'Ruxsat yo‘q'}, status=403)
    questions = Question.objects.filter(subject_id=subject_id)
    data = {
        'questions': [
            {
                'id': q.id,
                'text': q.text,
                'created_by': q.created_by.username if q.created_by else 'Noma’lum'
            }
            for q in questions
        ]
    }
    return JsonResponse(data)

# Testni tahrirlash (edit)
@login_required
def edit_test(request, test_id):
    login_redirect = login_check(request)
    if login_redirect:
        return login_redirect
    if not hasattr(request.user, 'role') or request.user.role != 'controller':
        return redirect('/api/login/')
    from main.models import TestQuestion
    test = Test.objects.get(id=test_id, created_by=request.user)
    groups = Group.objects.all()
    kafedralar = Kafedra.objects.all()
    bulimlar = Bulim.objects.all()
    subjects = Subject.objects.all()
    if request.method == 'POST':
        subject_id = request.POST.get('subject')
        question_count = int(request.POST.get('question_count'))
        total_score = int(request.POST.get('total_score'))
        duration_str = request.POST.get('duration')
        duration = parse_duration_string(duration_str)
        if duration is None:
            return render(request, 'controller_panel/edit_test.html', {
                'groups': groups,
                'kafedralar': kafedralar,
                'bulimlar': bulimlar,
                'subjects': subjects,
                'test': test,
                'error': 'Test muddati noto‘g‘ri formatda! To‘g‘ri format: soat:daqiq:soniya (masalan: 00:30:00)'
            })
        subject = Subject.objects.get(id=subject_id)
        # Test turi bo‘yicha mos maydonlarni yangilash
        if test.group:
            group_id = request.POST.get('group')
            group = Group.objects.get(id=group_id)
            test.group = group
            test.kafedra = None
            test.bulim = None
        elif test.kafedra:
            kafedra_id = request.POST.get('kafedra')
            kafedra = Kafedra.objects.get(id=kafedra_id)
            test.kafedra = kafedra
            test.group = None
            test.bulim = None
        elif test.bulim:
            bulim_id = request.POST.get('bulim')
            bulim = Bulim.objects.get(id=bulim_id)
            test.bulim = bulim
            test.group = None
            test.kafedra = None
        test.subject = subject
        test.question_count = question_count
        test.total_score = total_score
        test.duration = duration
        test.start_time = timezone.now()
        test.save()
        # Savollarni yangilash (oddiy variant: eski TestQuestionlarni o‘chirib, yangidan yaratamiz)
        TestQuestion.objects.filter(test=test).delete()
        questions = Question.objects.filter(subject=subject)
        if questions.count() < question_count:
            return render(request, 'controller_panel/edit_test.html', {
                'groups': groups,
                'kafedralar': kafedralar,
                'bulimlar': bulimlar,
                'subjects': subjects,
                'test': test,
                'error': 'Ushbu fanga yetarli savol mavjud emas!'
            })
        selected_questions = questions.order_by('?')[:question_count]
        score_per_question = total_score / question_count
        for question in selected_questions:
            TestQuestion.objects.create(test=test, question=question, score=score_per_question)
        return redirect('controller_dashboard')
    # GET
    return render(request, 'controller_panel/edit_test.html', {
        'groups': groups,
        'kafedralar': kafedralar,
        'bulimlar': bulimlar,
        'subjects': subjects,
        'test': test
    })

# Testni o‘chirish (delete)
@login_required
@csrf_exempt
def delete_test(request, test_id):
    login_redirect = login_check(request)
    if login_redirect:
        return login_redirect
    if not hasattr(request.user, 'role') or request.user.role != 'controller':
        return redirect('/api/login/')
    if request.method == 'POST':
        Test.objects.filter(id=test_id, created_by=request.user).delete()
    return redirect('controller_dashboard')


@login_required
def controller_dashboard(request):
    login_redirect = login_check(request)
    if login_redirect:
        return login_redirect
    if not hasattr(request.user, 'role') or request.user.role != 'controller':
        return redirect('/api/login/')
 
    tests = Test.objects.filter(created_by=request.user)
    now = timezone.now()
    # Har bir test uchun status va fanning nomini aniqlash
    test_list = []
    for test in tests:
        status = "Faol"
        if test.end_time and now > test.end_time:
            status = "Test muddati tugagan"
        subject_name = test.subject.name if test.subject else ""
        test_list.append({
            'test': test,
            'status': status,
            'subject_name': subject_name,
        })
    # frontendda tests deb ishlatiladi
    return render(request, 'controller_panel/dashboard.html', {'tests': test_list})








from datetime import timedelta

def parse_duration_string(duration_str):
    """Masalan: '00:30:00' ni timedelta ga aylantiradi"""
    try:
        h, m, s = map(int, duration_str.split(":"))
        return timedelta(hours=h, minutes=m, seconds=s)
    except:
        return None


from main.models import Kafedra, Bulim

@login_required
def add_test(request):
    login_redirect = login_check(request)
    if login_redirect:
        return login_redirect
    if not hasattr(request.user, 'role') or request.user.role != 'controller':
        return redirect('/api/login/')

    target = request.GET.get('target', 'student')
    context = {}
    if target == 'student':
        context['groups'] = Group.objects.all()
        group_id = request.GET.get('group') or request.POST.get('group')
        if group_id:
            from main.models import GroupSubject
            group_subjects = GroupSubject.objects.filter(group_id=group_id)
            context['subjects'] = Subject.objects.filter(id__in=group_subjects.values_list('subject_id', flat=True))
        else:
            context['subjects'] = Subject.objects.none()
        context['target'] = 'student'
    elif target == 'tutor':
        context['kafedralar'] = Kafedra.objects.all()
        context['subjects'] = Subject.objects.all()
        context['target'] = 'tutor'
    elif target == 'employee':
        context['bulimlar'] = Bulim.objects.all()
        context['subjects'] = Subject.objects.all()
        context['target'] = 'employee'
    else:
        context['groups'] = Group.objects.all()
        context['subjects'] = Subject.objects.all()
        context['target'] = 'student'

    if request.method == 'POST':
        subject_id = request.POST.get('subject')
        question_count = request.POST.get('question_count')
        total_score = request.POST.get('total_score')
        duration_str = request.POST.get('duration')
        minutes = request.POST.get('minutes', 30)
        # Majburiy maydonlar to'ldirilganini tekshirish
        if not (subject_id and question_count and total_score and duration_str):
            context['error'] = "Iltimos, barcha maydonlarni to'ldiring!"
            return render(request, 'controller_panel/add_test.html', context)
        try:
            question_count = int(question_count)
            total_score = int(total_score)
            minutes = int(minutes)
        except ValueError:
            context['error'] = "Sonli maydonlarga faqat raqam kiriting!"
            return render(request, 'controller_panel/add_test.html', context)
        duration = parse_duration_string(duration_str)
        context['error'] = None
        group = None
        kafedra = None
        bulim = None
        if target == 'student':
            group_id = request.POST.get('group')
            group = Group.objects.get(id=group_id)
        elif target == 'tutor':
            kafedra_id = request.POST.get('kafedra')
            kafedra = Kafedra.objects.get(id=kafedra_id)
        elif target == 'employee':
            bulim_id = request.POST.get('bulim')
            bulim = Bulim.objects.get(id=bulim_id)

        if duration is None:
            return render(request, 'controller_panel/add_test.html', context | {'error': 'Test muddati noto‘g‘ri formatda! To‘g‘ri format: soat:daqiq:soniya (masalan: 00:30:00)'})

        subject = Subject.objects.get(id=subject_id)
        from main.models import TestQuestion, Test
        # Test yaratish: modelga mos ravishda group, kafedra yoki bulimni saqlash kerak
        if target == 'student':
            test = Test.objects.create(
                group=group,
                subject=subject,
                question_count=question_count,
                total_score=total_score,
                duration=duration,
                minutes=minutes,
                created_by=request.user
            )
        elif target == 'tutor':
            # group o‘rniga kafedra saqlash uchun modelda mos o‘zgartirish kerak bo‘ladi
            test = Test.objects.create(
                group=None,
                subject=subject,
                question_count=question_count,
                total_score=total_score,
                duration=duration,
                minutes=minutes,
                created_by=request.user,
                kafedra=kafedra
            )
        elif target == 'employee':
            test = Test.objects.create(
                group=None,
                subject=subject,
                question_count=question_count,
                total_score=total_score,
                duration=duration,
                minutes=minutes,
                created_by=request.user,
                bulim=bulim
            )
        else:
            test = Test.objects.create(
                group=group,
                subject=subject,
                question_count=question_count,
                total_score=total_score,
                duration=duration,
                minutes=minutes,
                created_by=request.user
            )

        # Savollarni tanlash
        questions = Question.objects.filter(subject=subject)
        if questions.count() < question_count:
            context['error'] = 'Ushbu fanga yetarli savol mavjud emas!'
            return render(request, 'controller_panel/add_test.html', context)
        selected_questions = questions.order_by('?')[:question_count]
        score_per_question = total_score / question_count
        for question in selected_questions:
            TestQuestion.objects.create(test=test, question=question, score=score_per_question)
        return redirect('controller_dashboard')

    return render(request, 'controller_panel/add_test.html', context)
